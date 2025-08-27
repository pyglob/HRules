# scanner.py
import io
import os
import re
import cssutils
import pytesseract
from pathlib import Path
from typing import Dict, List, Tuple, Any
from PIL import Image, ImageOps
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR
from bs4 import BeautifulSoup
from pathlib import Path
from typing import Dict, List
import fitz  # PyMuPDF
from hrules.color_utils import contrast_ratio, CONTRAST_THRESHOLD, THEME_MAP, resolve_run_fg_hex

try:
    from psd_tools import PSDImage
    PSD_AVAILABLE = True
except ModuleNotFoundError:
    PSD_AVAILABLE = False

H_START = "<<<HIGHLIGHT>>>"
H_END = "<<<END>>>"
ZERO_WIDTH_CHARS = r'[\u200B\u200C\u200D\u2060\uFEFF]'
ZW_LABEL = "⟦ZW⟧"

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp"}
TEXT_EXTS = {".txt", ".css"}
HTML_EXTS = {".html", ".htm"}
DOC_EXTS = {".pdf", ".docx"}
PSD_EXTS = {".psd"}
AI_EXTS = {".ai"}  # PDF-compatible in many cases

EXIF_KEYS_OF_INTEREST = {
    270: "ImageDescription",
    315: "Artist",
    40092: "XPTitle",
    40094: "XPComment",
    40095: "XPAuthor",
    40096: "XPKeywords",
    40091: "XPSubject",
    37510: "UserComment"
}

HIDDEN_CSS_PATTERNS = [
    r'display\s*:\s*none',
    r'visibility\s*:\s*hidden',
    r'opacity\s*:\s*0(\.0+)?',
    r'opacity\s*:\s*0\.\d+'
]

INLINE_STYLE_COLOR_PAIR = re.compile(
    r'color\s*:\s*(#[0-9a-fA-F]{3,6}).*?background-color\s*:\s*(#[0-9a-fA-F]{3,6})',
    re.IGNORECASE | re.DOTALL
)


def hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
    h = hex_color.strip().lstrip('#')
    if len(h) == 3:
        h = ''.join(c*2 for c in h)
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def luminance(rgb: Tuple[int, int, int]) -> float:
    def channel(c):
        c = c / 255.0
        return c/12.92 if c <= 0.04045 else ((c+0.055)/1.055) ** 2.4
    r, g, b = [channel(v) for v in rgb]
    return 0.2126*r + 0.7152*g + 0.0722*b

# def contrast_ratio(fg_hex, bg_hex):
#     def hex_to_rgb(hex_color):
#         hex_color = hex_color.lstrip('#')
#         # Expand short form like #fff to #ffffff
#         if len(hex_color) == 3:
#             hex_color = ''.join([c*2 for c in hex_color])
#         return tuple(int(hex_color[i:i+2], 16) / 255.0 for i in (0, 2, 4))
#
#     def rel_luminance(rgb):
#         def channel(c):
#             return c / 12.92 if c <= 0.03928 else ((c + 0.055) / 1.055) ** 2.4
#         r, g, b = rgb
#         return 0.2126 * channel(r) + 0.7152 * channel(g) + 0.0722 * channel(b)
#
#     fg_rgb = hex_to_rgb(fg_hex)
#     bg_rgb = hex_to_rgb(bg_hex)
#     L1 = rel_luminance(fg_rgb)
#     L2 = rel_luminance(bg_rgb)
#     #ratio = (max(L1, L2) + 0.05) / (min(L1, L2) + 0.05)
#     return round((max(L1, L2) + 0.05) / (min(L1, L2) + 0.05), 2)


def detect_transparency(image_path_or_bytes) -> Tuple[bool, float]:
    if isinstance(image_path_or_bytes, (str, Path)):
        img = Image.open(image_path_or_bytes)
    else:
        img = Image.open(io.BytesIO(image_path_or_bytes))
    if img.mode != "RGBA":
        img = img.convert("RGBA")
    pixels = img.getdata()
    non_opaque = sum(1 for p in pixels if p[3] < 255)
    ratio = non_opaque / max(1, len(pixels))
    return ratio > 0, ratio


def ocr_image_bytes(image_bytes: bytes) -> str:
    try:
        img = Image.open(io.BytesIO(image_bytes))
        img = ImageOps.grayscale(img)
        text = pytesseract.image_to_string(img, config="--psm 6").strip()
        return text
    except Exception:
        return ""


def ocr_image_path(path: Path) -> str:
    try:
        img = Image.open(path)
        img = ImageOps.grayscale(img)
        text = pytesseract.image_to_string(img, config="--psm 6").strip()
        return text
    except Exception:
        return ""


def detect_hidden_chars(text: str) -> Tuple[int, str]:
    matches = list(re.finditer(ZERO_WIDTH_CHARS, text))
    highlighted = re.sub(ZERO_WIDTH_CHARS, ZW_LABEL, text)
    return len(matches), highlighted


def detect_low_contrast_in_text_blob(text: str) -> List[Dict[str, Any]]:
    findings = []
    for m in INLINE_STYLE_COLOR_PAIR.finditer(text):
        fg, bg = m.group(1), m.group(2)
        try:
            ratio = contrast_ratio(hex_to_rgb(fg), hex_to_rgb(bg))
        except Exception:
            continue
        if ratio < CONTRAST_THRESHOLD:
            snippet = text[max(0, m.start()-60):min(len(text), m.end()+60)]
            snippet = snippet.replace(m.group(0), f"{H_START}{m.group(0)}{H_END}")
            findings.append({"ratio": ratio, "snippet": snippet, "fg": fg, "bg": bg})
    return findings


def detect_hidden_css(text: str) -> List[str]:
    snippets = []
    for pat in HIDDEN_CSS_PATTERNS:
        for m in re.finditer(pat, text, flags=re.IGNORECASE):
            seg = text[m.start():m.end()]
            snippet = text[max(0, m.start()-60):min(len(text), m.end()+60)]
            snippets.append(snippet.replace(seg, f"{H_START}{seg}{H_END}"))
    return snippets


def analyze_css(css_text: str, selector_source="CSS") -> List[Dict[str, Any]]:
    out = []
    try:
        sheet = cssutils.parseString(css_text)
    except Exception:
        return out
    for rule in sheet:
        if rule.type != rule.STYLE_RULE:
            continue
        color = None
        bg = None
        has_hidden = False
        hidden_snippets = []
        for prop in rule.style:
            name = prop.name.lower()
            val = prop.value.strip()
            if name == "color" and val.startswith("#"):
                color = val
            elif name == "background-color" and val.startswith("#"):
                bg = val
            elif name in ("display", "visibility", "opacity"):
                if re.search("|".join(HIDDEN_CSS_PATTERNS), f"{name}:{val}", flags=re.IGNORECASE):
                    has_hidden = True
                    hidden_snippets.append(f"{rule.selectorText} {{ {H_START}{name}:{val}{H_END} }}")
        if has_hidden:
            out.append({"type": "hidden_css", "selector": rule.selectorText, "snippet": "\n".join(hidden_snippets)})
        if color and bg:
            try:
                ratio = contrast_ratio(hex_to_rgb(color), hex_to_rgb(bg))
            except Exception:
                continue
            if ratio < CONTRAST_THRESHOLD:
                snippet = f"{rule.selectorText} {{ {H_START}color:{color}; background-color:{bg};{H_END} }}"
                out.append({"type": "low_contrast", "selector": rule.selectorText, "ratio": ratio, "snippet": snippet})
    return out


def scan_exif(path: Path) -> List[str]:
    findings = []
    try:
        img = Image.open(path)
        exif = img.getexif()
        if not exif:
            return findings
        for tag_id, value in exif.items():
            name = EXIF_KEYS_OF_INTEREST.get(tag_id)
            if not name:
                continue
            if isinstance(value, bytes):
                try:
                    value = value.decode("utf-16le", errors="ignore")
                except Exception:
                    value = value.decode("utf-8", errors="ignore")
            findings.append(f"{name}: {str(value)[:200]}")
    except Exception:
        pass
    return findings


def scan_pdf(path: Path) -> Dict[str, List[str]]:
    v, n = [], []
    doc = fitz.open(str(path))

    # --- Low-contrast text detection ---
    for page_num, page in enumerate(doc, start=1):
        try:
            text_dict = page.get_text("dict")
            for block in text_dict.get("blocks", []):
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        try:
                            color_val = span.get("color")
                            fg_hex = None

                            # Handle tuple of floats (0..1)
                            if isinstance(color_val, tuple) and len(color_val) == 3:
                                r, g, b = [int(c * 255) for c in color_val]
                                fg_hex = f"#{r:02x}{g:02x}{b:02x}"

                            # Handle integer color (e.g., 16777215 for white)
                            elif isinstance(color_val, int):
                                r = (color_val >> 16) & 255
                                g = (color_val >> 8) & 255
                                b = color_val & 255
                                fg_hex = f"#{r:02x}{g:02x}{b:02x}"

                            if fg_hex:
                                bg_hex = "#FFFFFF"  # assume white background
                                ratio = contrast_ratio(fg_hex, bg_hex)
                                if ratio < CONTRAST_THRESHOLD:
                                    v.append(f"PDF low-contrast text on page {page_num}: {fg_hex} on {bg_hex} ")
                                    if span["text"].strip():
                                        n.append(f"Excerpt (p{page_num}): {span['text']}")
                        except Exception as e:
                            print(f"ERROR in span loop: {e}")
        except Exception:
            pass

    # --- Hidden/zero-width characters ---
    full_text = "\n".join(page.get_text() for page in doc)
    hidden_count, highlighted = detect_hidden_chars(full_text)
    if hidden_count:
        v.append(f"PDF hidden/zero-width text: {hidden_count} occurrences ")
        n.append("PDF text excerpt:\n" + highlighted[:800] + ("\n..." if len(highlighted) > 800 else ""))

    # --- Images: transparency + OCR ---
    for page_num, page in enumerate(doc, start=1):
        images = page.get_images(full=True)
        for img_index, img in enumerate(images):
            try:
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                img_bytes = pix.tobytes("png")
                has_trans, ratio = detect_transparency(img_bytes)
                if has_trans:
                    v.append(f"PDF page {page_num} image transparency: {ratio*100:.2f}% ")
                ocr = ocr_image_bytes(img_bytes)
                if ocr:
                    pdf_ocr_excerpt = ocr[:300].replace("\n", " ")
                    n.append(f"PDF page {page_num} image OCR text: {pdf_ocr_excerpt}")
            except Exception:
                pass

    return {"violations": v, "notes": n}


def scan_docx(path: Path) -> Dict[str, List[str]]:
    v, n = [], []
    doc = Document(str(path))
    full_text = []
    hidden_runs = 0
    seen_excerpts = set()  # prevent duplicate entries

    def iter_all_paragraphs(doc):
        # Body paragraphs
        for p in doc.paragraphs:
            yield p
        # Tables in body
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
        # Headers/footers + any tables inside them
        for section in doc.sections:
            for p in section.header.paragraphs:
                yield p
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p
            for p in section.footer.paragraphs:
                yield p
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:  # fixed: iterate paragraphs, not undefined p
                            yield p

    for para in iter_all_paragraphs(doc):
        full_text.append(para.text)
        for run in para.runs:
            t = run.text
            fg_hex = resolve_run_fg_hex(run, THEME_MAP)
            if not fg_hex:
                fg_hex = THEME_MAP[MSO_THEME_COLOR.TEXT_1]  # default to black

            # Count hidden runs
            try:
                if run.font.hidden:
                    hidden_runs += 1
            except Exception:
                pass

            text = t.strip()
            if not text:
                continue

            # Contrast check
            ratio = contrast_ratio(fg_hex, "#ffffff")
            if ratio < CONTRAST_THRESHOLD:
                key = (fg_hex, text)
                if key not in seen_excerpts:
                    seen_excerpts.add(key)
                    v.append(f"DOCX low-contrast text: {fg_hex} on #ffffff")
                    n.append(f"Excerpt: {text}")

    if hidden_runs:
        v.append(f"DOCX hidden text runs: {hidden_runs} ")

    # hidden/zero-width characters
    text_all = "\n".join(full_text)
    hidden_count, highlighted = detect_hidden_chars(text_all)
    if hidden_count:
        v.append(f"DOCX hidden/zero-width text: {hidden_count} occurrences ")
        n.append("DOCX text excerpt:\n" + highlighted[:800] + ("...\n" if len(highlighted) > 800 else ""))

    # images: transparency + OCR
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img_bytes = rel.target_part.blob
            has_trans, ratio = detect_transparency(img_bytes)
            if has_trans:
                v.append(f"DOCX image transparency: {ratio*100:.2f}% ")
            ocr = ocr_image_bytes(img_bytes)
            if ocr:
                cleaned_ocr = ocr[:300].replace("\n", " ")
                n.append("DOCX image OCR text: " + cleaned_ocr)

    return {"violations": v, "notes": n}


def scan_html(path: Path) -> Dict[str, List[str]]:
    v, n = [], []
    html = path.read_text(encoding="utf-8", errors="ignore")
    count, highlighted = detect_hidden_chars(html)
    if count:
        v.append(f"HTML hidden/zero-width characters: {count} ")
        n.append("HTML excerpt:\n" + highlighted[:800] + ("\n..." if len(highlighted) > 800 else ""))
    # Inline styles
    for item in detect_low_contrast_in_text_blob(html):
        v.append(f"Low-contrast inline style (ratio {item['ratio']:.2f}) ")
        n.append("Inline snippet:\n" + item["snippet"])
    for snip in detect_hidden_css(html):
        v.append("Hidden CSS detected (inline) ")
        n.append(snip)
    # Embedded and linked CSS
    soup = BeautifulSoup(html, "html.parser")
    for style in soup.find_all("style"):
        findings = analyze_css(style.get_text(), "embedded CSS")
        for it in findings:
            if it["type"] == "low_contrast":
                v.append(f"Low-contrast CSS selector {it['selector']} (ratio {it['ratio']:.2f}) ")
                n.append(it["snippet"])
            elif it["type"] == "hidden_css":
                v.append(f"Hidden CSS {it['selector']} ")
                n.append(it["snippet"])
    for link in soup.find_all("link", rel=lambda v: v and "stylesheet" in v):
        href = link.get("href")
        if not href:
            continue
        css_path = (path.parent / href).resolve()
        if css_path.exists() and css_path.is_file():
            css_text = css_path.read_text(encoding="utf-8", errors="ignore")
            findings = analyze_css(css_text, href)
            for it in findings:
                if it["type"] == "low_contrast":
                    v.append(f"Low-contrast CSS {it['selector']} (ratio {it['ratio']:.2f}) in {href} ")
                    n.append(it["snippet"])
                elif it["type"] == "hidden_css":
                    v.append(f"Hidden CSS {it['selector']} in {href} ")
                    n.append(it["snippet"])
    return {"violations": v, "notes": n}


def scan_psd(path: Path) -> Dict[str, List[str]]:
    v, n = [], []
    if not PSD_AVAILABLE:
        n.append("psd-tools not installed; PSD layer scan skipped.")
        return {"violations": v, "notes": n}
    try:
        psd = PSDImage.open(str(path))
    except Exception as e:
        n.append(f"PSD parse failed: {e}")
        return {"violations": v, "notes": n}
    def walk(layer, trail=""):
        name = getattr(layer, "name", "unnamed")
        visible = getattr(layer, "visible", True)
        opacity = getattr(layer, "opacity", 255)
        label = f"{trail}/{name}".strip("/")
        if not visible:
            v.append(f"PSD hidden layer: {label} ")
        if opacity < 255:
            v.append(f"PSD semi-transparent layer: {label} (opacity {opacity}/255) ")
        if hasattr(layer, "layers"):
            for child in layer.layers:
                walk(child, label)
    walk(psd)
    return {"violations": v, "notes": n}


def scan_image(path: Path) -> Dict[str, List[str]]:
    v, n = [], []
    has_trans, ratio = detect_transparency(path)
    if has_trans:
        v.append(f"Image transparency: {ratio*100:.2f}% ")
    exif = scan_exif(path)
    for line in exif:
        n.append(f"EXIF {line}")
    ocr = ocr_image_path(path)
    if ocr:
        image_ocr_excerpt = ocr[:300].replace("\n", " ")
        n.append(f"Image OCR text: {image_ocr_excerpt}")
        # n.append(f"DOCX image OCR text: {ocr[:300].replace('\n', ' ')}")  issue in python 3.10
    return {"violations": v, "notes": n}


def scan_text_or_css(path: Path) -> Dict[str, List[str]]:
    v, n = [], []
    text = path.read_text(encoding="utf-8", errors="ignore")
    count, highlighted = detect_hidden_chars(text)
    if count:
        v.append(f"{path.suffix.upper()} hidden/zero-width text: {count} ")
        n.append("Excerpt:\n" + highlighted[:800] + ("\n..." if len(highlighted) > 800 else ""))
    if path.suffix.lower() == ".css":
        findings = analyze_css(text, str(path))
        for it in findings:
            if it["type"] == "low_contrast":
                v.append(f"Low-contrast CSS {it['selector']} (ratio {it['ratio']:.2f}) ")
                n.append(it["snippet"])
            elif it["type"] == "hidden_css":
                v.append(f"Hidden CSS {it['selector']} ")
                n.append(it["snippet"])
    return {"violations": v, "notes": n}


def scan_file(path: Path) -> Dict[str, List[str]]:
    #print(f"DEBUG: scan_file called for {path} with ext={path.suffix.lower()}")
    ext = path.suffix.lower().strip()
    if ext in IMAGE_EXTS:
        return scan_image(path)
    if ext in TEXT_EXTS:
        return scan_text_or_css(path)
    if ext in HTML_EXTS:
        return scan_html(path)
    if ext == ".pdf":
        return scan_pdf(path)
    if ext == ".docx":
        return scan_docx(path)
    if ext in PSD_EXTS:
        return scan_psd(path)
    if ext in AI_EXTS:
        try:
            with open(path, "rb") as f:
                header = f.read(4)
            if header.startswith(b"%PDF"):
                return scan_pdf(path)
        except Exception:
            pass
        return {"violations": [], "notes": ["AI file not PDF-compatible; deep scan skipped."]}
    return {"violations": [], "notes": [f"Unsupported file type: {ext}"]}


def scan_directory(dir_path: Path) -> List[Tuple[Path, Dict[str, List[str]]]]:
    #print(f"DEBUG: scan_directory called for {dir_path} with ext={dir_path.suffix.lower()}")
    results = []
    for root, _, files in os.walk(dir_path):
        for name in files:
            fp = Path(root) / name
            results.append((fp, scan_file(fp)))
    return results
